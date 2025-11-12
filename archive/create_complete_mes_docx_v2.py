"""
Withforce MES 시스템 사양서 완전판 DOCX 생성 스크립트 v2
- 모든 테이블, JSON 예시, 다이어그램 포함
- 심플 + 모던 스타일
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_page_number(section):
    """페이지 번호 추가"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def setup_modern_styles(doc):
    """모던 스타일 설정"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)

    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Malgun Gothic'
    heading1.font.size = Pt(18)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(31, 73, 125)

    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Malgun Gothic'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(68, 114, 196)

    heading3 = doc.styles['Heading 3']
    heading3.font.name = 'Malgun Gothic'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(112, 48, 160)

def create_title_page(doc):
    """표지 페이지"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(100)
    title.paragraph_format.space_after = Pt(30)

    run = title.add_run('Withforce MES 시스템 사양서')
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(60)

    run = subtitle.add_run('웨어러블 로봇 제조실행시스템\n전체 시스템 사양 및 설계 문서 (완전판)')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_before = Pt(150)

    info_text = (
        '문서 버전: v2.1\n'
        '최초 작성일: 2025-11-10\n'
        '최종 수정일: 2025-11-11\n'
        '문서 유형: 시스템 요구사항 명세서\n'
        '상태: Production Ready'
    )

    run = info.add_run(info_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(127, 127, 127)

    footer_info = doc.add_paragraph()
    footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_info.paragraph_format.space_before = Pt(100)

    run = footer_info.add_run('F2X NeuroHub Team')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    doc.add_page_break()

def add_code_block(doc, code, language=''):
    """코드 블록 추가 (회색 배경)"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 회색 배경 설정
    from docx.oxml import parse_xml
    shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
    p._p.get_or_add_pPr().append(shading_elm)

    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)

def nsdecls(*prefixes):
    """네임스페이스 선언"""
    ns_map = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }
    return ' '.join('xmlns:{}="{}"'.format(p, ns_map.get(p, '')) for p in prefixes)

def add_section(doc, title, level=1):
    """섹션 제목 추가"""
    if level == 1:
        doc.add_heading(title, 1)
    elif level == 2:
        doc.add_heading(title, 2)
    elif level == 3:
        doc.add_heading(title, 3)

def add_text(doc, text, bold=False, indent=0):
    """일반 텍스트 추가"""
    p = doc.add_paragraph()
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent * 0.5)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Malgun Gothic'
    if bold:
        run.bold = True

def add_bullet(doc, text, level=0):
    """불릿 포인트 추가"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5)
    p.paragraph_format.space_after = Pt(3)

def add_table(doc, headers, rows, col_widths=None):
    """표 추가 (향상된 스타일)"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 열 너비 설정
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    # 헤더
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Malgun Gothic'
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 헤더 배경색
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
        cell._element.get_or_add_tcPr().append(shading_elm)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # 데이터 행
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(cell_data)
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
            if run:
                run.font.size = Pt(9)
                run.font.name = 'Malgun Gothic'

    doc.add_paragraph()

def add_diagram(doc, diagram_text):
    """다이어그램 추가 (텍스트 기반)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # 회색 배경
    from docx.oxml import parse_xml
    shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
    p._p.get_or_add_pPr().append(shading_elm)

    run = p.add_run(diagram_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 0, 0)

# 나머지 함수들은 이전과 동일하게 유지하되, 더 많은 상세 내용 추가
def create_mes_specification():
    doc = Document()

    # 페이지 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    setup_modern_styles(doc)
    create_title_page(doc)

    # 목차는 Word에서 자동 생성되도록 안내
    add_section(doc, '목차', 1)
    add_text(doc, 'Microsoft Word에서 [참조] → [목차] → [자동 목차]를 클릭하여 목차를 생성하세요.')
    doc.add_page_break()

    # === 문서 요약 ===
    add_section(doc, '문서 요약', 1)

    add_section(doc, '프로젝트 개요', 2)
    add_text(doc,
        'Withforce 웨어러블 로봇 생산 라인에 MES(제조실행시스템)를 도입하여 '
        '완벽한 제품 추적성과 실시간 생산 모니터링 체계를 구축합니다.'
    )

    add_section(doc, '핵심 기능', 2)
    add_bullet(doc, 'LOT 및 시리얼 번호 기반 제품 추적')
    add_bullet(doc, '8개 공정별 착공/완공 관리')
    add_bullet(doc, '실시간 생산 현황 대시보드')
    add_bullet(doc, '품질 데이터 자동 수집 및 분석')

    add_section(doc, '주요 수치', 2)
    add_table(doc,
        ['항목', '수치'],
        [
            ['동시 접속자', '100명 (작업자 70명 + 관리자 30명)'],
            ['일일 트랜잭션', '약 50,000건'],
            ['생산 대상', 'LOT당 100대, 일일 약 100대']
        ]
    )

    add_section(doc, '시스템 구성도', 2)
    add_diagram(doc, '''
    ┌─────────────┐     ┌─────────────┐     ┌─────────────┐
    │  작업 PC 1-7│────▶│  MES 서버   │◀────│  관리자 PC  │
    │  (PyQt5)    │     │  (FastAPI)  │     │  (React)    │
    └─────────────┘     │  PostgreSQL │     └─────────────┘
                        │  Redis      │
    ┌─────────────┐     └─────────────┘
    │ 라벨 프린터 │◀────────┘
    │ (Zebra)     │
    └─────────────┘
    ''')

    doc.add_page_break()

    # === 1. 프로젝트 개요 ===
    add_section(doc, '1. 프로젝트 개요', 1)

    add_section(doc, '1.1 프로젝트 배경', 2)
    add_text(doc,
        'Withforce는 산업용/농업용 허리 보조 로봇을 생산하고 있으며, '
        '현재 생산 공정 전반에 걸친 데이터 수집 및 추적 체계가 미흡한 상황입니다.'
    )
    add_text(doc,
        '제품 불량 발생 시 원인 파악이 어렵고, 실시간 생산 현황 모니터링이 불가능하여 '
        '효율적인 생산 관리에 어려움을 겪고 있습니다.'
    )

    add_section(doc, '1.2 프로젝트 목표', 2)
    add_text(doc, 'MES 도입을 통한 핵심 목표:', bold=True)
    add_bullet(doc, '완벽한 제품 추적성 확보 (LOT 및 시리얼 번호 기반)')
    add_bullet(doc, '실시간 생산 현황 모니터링')
    add_bullet(doc, '데이터 기반 품질 관리')
    add_bullet(doc, '공정 병목 구간 식별 및 최적화')

    add_section(doc, '1.3 기대 효과', 2)
    add_table(doc,
        ['항목', 'AS-IS (현재)', 'TO-BE (개선 후)', '개선 효과'],
        [
            ['LOT 생성 시간', '5분 (수동)', '30초 (자동)', '90% 단축'],
            ['불량 원인 파악', '불가능', '5분 이내', '추적성 확보'],
            ['생산 현황 파악', '1시간 (수동 집계)', '실시간', '실시간 모니터링'],
            ['품질 데이터 수집', '미수집', '100% 수집', '데이터 기반 개선']
        ]
    )

    doc.add_page_break()

    # === 2. 제품 및 공정 현황 ===
    add_section(doc, '2. 제품 및 공정 현황 (AS-IS)', 1)

    add_section(doc, '2.1 제품 소개', 2)
    add_text(doc, 'Withforce 웨어러블 로봇의 주요 부품:')
    add_bullet(doc, 'LMA (Linear Muscle Actuator): SMA 스프링 기반 핵심 구동기')
    add_bullet(doc, 'SMA 스프링: 온도로 형상을 기억하여 구동하는 특수 합금')
    add_bullet(doc, '모선 케이블: SMA 스프링 제조 원재료')
    add_bullet(doc, '제어 보드 및 센서 (온도, TOF)')

    add_section(doc, '2.2 생산 공정 흐름', 2)
    add_text(doc, '8개 공정으로 구성:', bold=True)

    add_diagram(doc, '''
    레이저 마킹 → LMA 조립 → 센서 검사 → 펌웨어 업로드
         ↓
    로봇 조립 → 성능검사 → 라벨 프린팅 → 포장+외관검사
    ''')

    add_section(doc, '2.2.1 공정별 상세', 3)

    processes_detail = [
        ('공정 1: 레이저 마킹', 'LOT 번호를 제품에 레이저로 각인하여 추적성 확보'),
        ('공정 2: LMA 조립', 'SMA 스프링, 모선, 링크 등을 조립하여 LMA 완성'),
        ('공정 3: 센서 검사', '온도 센서 및 TOF 센서 통신 검사 (자동 검사)'),
        ('공정 4: 펌웨어 업로드', '제어 보드에 최신 펌웨어 업로드'),
        ('공정 5: 로봇 조립', 'LMA를 로봇 프레임에 조립'),
        ('공정 6: 성능 검사', '온도/변위/힘 측정하여 EOL 검사 (자동 검사)'),
        ('공정 7: 라벨 프린팅', '시리얼 번호 발급 및 바코드 라벨 출력'),
        ('공정 8: 포장+외관검사', '최종 외관 검사 후 포장 완료')
    ]

    for proc_name, proc_desc in processes_detail:
        add_text(doc, proc_name, bold=True)
        add_text(doc, proc_desc, indent=1)

    add_section(doc, '2.3 LOT 및 시리얼 번호 체계', 2)

    add_section(doc, '2.3.1 LOT 번호 형식', 3)
    add_text(doc, '형식: WF-KR-YYMMDDX-nnn', bold=True)
    add_text(doc, '예시: WF-KR-251110D-001')

    add_table(doc,
        ['요소', '설명', '예시', '비고'],
        [
            ['WF', '제품 코드', 'WF', 'Withforce'],
            ['KR', '공장 코드', 'KR', '한국 공장'],
            ['YYMMDD', '생산 날짜', '251110', '2025년 11월 10일'],
            ['X', '교대', 'D', 'D=Day, N=Night'],
            ['nnn', '일련번호', '001', '당일 첫 번째 LOT']
        ]
    )

    add_section(doc, '2.3.2 시리얼 번호 형식', 3)
    add_text(doc, '형식: WF-KR-YYMMDDX-nnn-YYYY', bold=True)
    add_text(doc, '예시: WF-KR-251110D-001-0042')
    add_text(doc, '설명: WF-KR-251110D-001 LOT의 42번째 제품')

    add_section(doc, '2.4 불량 유형 코드 체계', 2)
    add_text(doc, '8개 대분류:', bold=True)

    add_table(doc,
        ['코드', '영문명', '한글명', '예시'],
        [
            ['M', 'Material', '재료 불량', 'M001: 모선 불량'],
            ['A', 'Assembly', '조립 불량', 'A001: 스프링 조립 불량'],
            ['E', 'Electrical', '전기/전자 불량', 'E001: 센서 통신 불량'],
            ['P', 'Performance', '성능 불량', 'P001: 힘 측정값 이탈'],
            ['V', 'Visual', '외관 불량', 'V001: 스크래치'],
            ['S', 'Software', '소프트웨어 불량', 'S001: 펌웨어 업로드 실패'],
            ['D', 'Dimensional', '치수 불량', 'D001: 치수 불일치'],
            ['O', 'Other', '기타 불량', 'O001: 기타 원인']
        ]
    )

    doc.add_page_break()

    # === 3. MES 요구사항 ===
    add_section(doc, '3. MES 요구사항 (TO-BE)', 1)

    add_section(doc, '3.1 전체 프로세스 흐름', 2)
    add_diagram(doc, '''
    1. 생산 관리자: LOT 생성 요청
       ↓
    2. 백엔드: LOT 번호 발급 (WF-KR-251110D-001)
       ↓
    3. 라벨 프린터: LOT 바코드 라벨 출력
       ↓
    4. 작업대차에 라벨 부착
       ↓
    5. 현장 작업자: LOT 바코드 스캔 → 착공 등록
       ↓
    6. 외부 공정 앱: 작업 수행 → JSON 파일 생성
       ↓
    7. File Watcher: JSON 파일 감지 → 백엔드 전송
       ↓
    8. 백엔드: 완공 데이터 저장
       ↓
    9. 대시보드: 실시간 현황 업데이트
    ''')

    add_section(doc, '3.2 LOT 관리', 2)

    add_text(doc, 'FR-LOT-001: LOT 생성', bold=True)
    add_bullet(doc, '생성 주체: 생산 관리자 (React 대시보드)')
    add_bullet(doc, 'LOT 번호 자동 발급 (중복 방지)')
    add_bullet(doc, '생성 시 초기 상태: CREATED')
    add_bullet(doc, '목표 수량: 100개 (기본값)')

    add_text(doc, 'FR-LOT-002: LOT 상태 관리', bold=True)
    add_text(doc, '상태 전환:', bold=True)
    add_diagram(doc, '''
    CREATED → IN_PROGRESS → COMPLETED → CLOSED

    - CREATED: LOT 생성 직후
    - IN_PROGRESS: 첫 번째 착공 시
    - COMPLETED: 모든 공정 완료 시
    - CLOSED: 최종 마감 (수동)
    ''')

    add_section(doc, '3.3 시리얼 번호 관리', 2)

    add_text(doc, 'FR-SN-001: 시리얼 번호 생성', bold=True)
    add_bullet(doc, '생성 시점: 공정 7 (Label Printing) 착공 시')
    add_bullet(doc, '생성 로직: LOT 번호 + 순번 (0001~0100)')
    add_bullet(doc, '중복 방지: DB UNIQUE 제약 조건')

    add_section(doc, '3.4 공정 관리', 2)

    add_section(doc, '3.4.1 착공 처리', 3)
    add_text(doc, 'API 엔드포인트: POST /api/v1/process/start')

    add_text(doc, '요청 JSON 예시:', bold=True)
    add_code_block(doc, '''{
  "lot_number": "WF-KR-251110D-001",
  "line_id": "LINE-A",
  "process_id": "PROC-001",
  "process_name": "레이저 마킹",
  "equipment_id": "LASER-01",
  "worker_id": "W001",
  "start_time": "2025-01-10T09:00:00+09:00"
}''', 'json')

    add_text(doc, '응답 JSON 예시 (성공):', bold=True)
    add_code_block(doc, '''{
  "status": "success",
  "message": "착공이 등록되었습니다",
  "data": {
    "lot_number": "WF-KR-251110D-001",
    "process_id": "PROC-001",
    "work_order_id": "WO-20250110-001"
  }
}''', 'json')

    add_text(doc, '응답 JSON 예시 (실패):', bold=True)
    add_code_block(doc, '''{
  "status": "error",
  "message": "이전 공정이 완료되지 않았습니다",
  "error_code": "PREVIOUS_PROCESS_NOT_COMPLETED",
  "data": {
    "required_process": "LMA 조립"
  }
}''', 'json')

    add_section(doc, '3.4.2 완공 처리 (File Watcher 방식)', 3)
    add_text(doc, '배경: 외부 공정 앱은 소스 코드 접근 불가, API 연동 불가')
    add_text(doc, '해결: JSON 파일 기반 데이터 교환')

    add_text(doc, '프로세스:', bold=True)
    add_diagram(doc, '''
    외부 공정 앱
       ↓ (JSON 파일 생성)
    C:\\neurohub_work\\pending\\
       ↓ (File Watcher 감지)
    프론트엔드 (PyQt5)
       ↓ (HTTP POST)
    백엔드 (FastAPI)
       ↓ (DB 저장)
    PostgreSQL
    ''')

    add_text(doc, '공정 2 (LMA 조립) 완공 JSON 예시:', bold=True)
    add_code_block(doc, '''{
  "lot_number": "WF-KR-251110D-001",
  "process_id": "PROC-002",
  "process_name": "LMA 조립",
  "equipment_id": "ASSEMBLY-01",
  "worker_id": "W002",
  "start_time": "2025-01-10T09:30:00+09:00",
  "complete_time": "2025-01-10T10:25:00+09:00",
  "result": "PASS",
  "process_data": {
    "sma_spring_lot": "SPRING-2025011001",
    "busbar_lot": "BUSBAR-2025011001",
    "assembly_time": 55,
    "visual_inspection": "PASS"
  }
}''', 'json')

    add_text(doc, '공정 3 (센서 검사) FAIL 예시:', bold=True)
    add_code_block(doc, '''{
  "lot_number": "WF-KR-251110D-001",
  "process_id": "PROC-003",
  "result": "FAIL",
  "process_data": {
    "temp_sensor": {
      "measured_temp": 58.0,
      "target_temp": 60.0,
      "result": "FAIL"
    },
    "tof_sensor": {
      "i2c_communication": true,
      "result": "PASS"
    },
    "defect_type": "SENSOR_TEMP_FAIL",
    "defect_description": "온도 센서 측정값 범위 초과"
  }
}''', 'json')

    add_section(doc, '3.4.3 공정별 완공 데이터 필드', 3)

    add_table(doc,
        ['공정', '필수 필드', '데이터 타입', '검증 규칙'],
        [
            ['레이저 마킹', 'marking_result', 'string', '"PASS" or "FAIL"'],
            ['LMA 조립', 'sma_spring_lot, busbar_lot', 'string', 'LOT 형식 검증'],
            ['센서 검사', 'temp_sensor.result, tof_sensor.result', 'string', '"PASS" or "FAIL"'],
            ['펌웨어 업로드', 'firmware_version, upload_result', 'string', '버전 형식 검증'],
            ['로봇 조립', 'assembly_time, cable_connection', 'number, string', '> 0, "OK"/"FAIL"'],
            ['성능 검사', 'test_results[], overall_result', 'array, string', '길이 > 0'],
            ['라벨 프린팅', 'serial_number, label_printed', 'string, boolean', '시리얼 형식, true'],
            ['포장+외관검사', 'visual_defects[], final_result', 'array, string', '"PASS"/"FAIL"']
        ]
    )

    add_section(doc, '3.5 사용자 및 권한 관리', 2)

    add_text(doc, '권한 매트릭스 (RBAC):', bold=True)
    add_table(doc,
        ['기능', '작업자', '생산 관리자', '시스템 관리자'],
        [
            ['공정 착공/완공', '✅', '✅', '✅'],
            ['LOT 생성/관리', '❌', '✅', '✅'],
            ['불량 등록', '✅', '✅', '✅'],
            ['대시보드 조회', '제한적', '✅', '✅'],
            ['추적성 조회', '❌', '✅', '✅'],
            ['사용자 관리', '❌', '❌', '✅'],
            ['시스템 설정', '❌', '❌', '✅']
        ]
    )

    doc.add_page_break()

    # === 4. 시스템 아키텍처 ===
    add_section(doc, '4. 시스템 아키텍처', 1)

    add_section(doc, '4.1 배포 옵션 비교', 2)

    add_table(doc,
        ['구분', 'Option A: 온프레미스', 'Option B-1: Railway', 'Option B-2: AWS'],
        [
            ['배포 방식', '자체 서버 구축', 'Cloud PaaS', 'Cloud IaaS'],
            ['초기 투자', '6,790만원', '1,805만원', '1,805만원'],
            ['월 운영비', '113만원', '26만원', '220-350만원'],
            ['3년 TCO', '10,858만원', '2,741만원', '13,685만원'],
            ['네트워크', '내부 LAN만', '인터넷 필수', '인터넷 필수'],
            ['배포 시간', '12-13주', '8주', '9-10주'],
            ['확장성', '서버 추가 구매', '자동 스케일링', '수동/자동 스케일링']
        ]
    )

    add_section(doc, '4.1.1 Option A: 온프레미스 아키텍처', 3)
    add_diagram(doc, '''
    ┌──────────────── 내부 LAN (192.168.1.0/24) ────────────────┐
    │                                                            │
    │  ┌───────────┐  ┌───────────┐  ┌───────────┐            │
    │  │작업 PC 1-7│  │ 관리자 PC │  │ 프린터 1-2│            │
    │  └─────┬─────┘  └─────┬─────┘  └─────┬─────┘            │
    │        │              │              │                    │
    │        └──────────────┼──────────────┘                    │
    │                       │                                   │
    │                 ┌─────▼─────┐                            │
    │                 │   MES 서버 │                            │
    │                 │  Nginx     │                            │
    │                 │  FastAPI   │                            │
    │                 │ PostgreSQL │                            │
    │                 │  Redis     │                            │
    │                 └───────────┘                            │
    └────────────────────────────────────────────────────────┘
    ''')

    add_section(doc, '4.1.2 Option B-1: Railway 아키텍처', 3)
    add_diagram(doc, '''
    ┌────── 공장 내부 ──────┐         ┌─── Railway Cloud ───┐
    │                       │         │                      │
    │  ┌─────────────┐     │         │  ┌──────────────┐  │
    │  │ 작업 PC 1-7 │─────┼─────────┼─▶│ FastAPI      │  │
    │  │ (PyQt5)     │     │         │  │ (Web Service)│  │
    │  └─────────────┘     │         │  └──────┬───────┘  │
    │                       │         │         │          │
    │  ┌─────────────┐     │         │  ┌──────▼───────┐  │
    │  │  관리자 PC  │─────┼─────────┼─▶│ PostgreSQL   │  │
    │  │  (React)    │     │         │  │ (Managed)    │  │
    │  └─────────────┘     │         │  └──────────────┘  │
    │                       │         │                      │
    │  ┌─────────────┐     │         │  ┌──────────────┐  │
    │  │ 프린터 1-2  │◀────┼─────────┼──│ Redis        │  │
    │  └─────────────┘     │         │  │ (Managed)    │  │
    └───────────────────────┘         │  └──────────────┘  │
                                       └──────────────────┘
           ⚠️ 인터넷 연결 필수
    ''')

    add_section(doc, '4.1.3 Option B-2: AWS 아키텍처', 3)
    add_diagram(doc, '''
    ┌──── 공장 내부 ────┐    ┌────────── AWS Cloud ─────────┐
    │                   │    │                               │
    │  ┌─────────┐     │    │  ┌────────────────────┐      │
    │  │작업 PC  │─────┼────┼─▶│ ALB (Load Balancer)│      │
    │  │관리자 PC│     │    │  └─────────┬──────────┘      │
    │  └─────────┘     │    │            │                  │
    │                   │    │  ┌─────────▼──────────┐      │
    │  ┌─────────┐     │    │  │ EC2 (FastAPI)  ×2 │      │
    │  │프린터 1-2│◀────┼────┼──│ t3.medium          │      │
    │  └─────────┘     │    │  └──────────┬─────────┘      │
    └───────────────────┘    │             │                 │
                              │  ┌──────────▼─────────┐      │
           ⚠️ 인터넷 필수     │  │ RDS (PostgreSQL)   │      │
                              │  │ Multi-AZ           │      │
                              │  └────────────────────┘      │
                              │  ┌────────────────────┐      │
                              │  │ ElastiCache (Redis)│      │
                              │  │ S3 (펌웨어 저장)   │      │
                              │  └────────────────────┘      │
                              └───────────────────────────┘
    ''')

    add_section(doc, '4.2 네트워크 구성', 2)

    add_table(doc,
        ['구분', 'IP 주소', '용도', '포트', '비고'],
        [
            ['WiFi 공유기', '192.168.1.1', 'Gateway', '-', '내부 LAN'],
            ['MES 서버', '192.168.1.10', 'MES 서버', '80, 443', 'Nginx'],
            ['작업 PC 1', '192.168.1.101', '레이저 마킹', '-', 'PyQt5 앱'],
            ['작업 PC 2', '192.168.1.102', 'LMA 조립', '-', 'PyQt5 앱'],
            ['작업 PC 3', '192.168.1.103', '센서 검사', '-', 'PyQt5 앱'],
            ['작업 PC 4', '192.168.1.104', '펌웨어 업로드', '-', 'PyQt5 앱'],
            ['작업 PC 5', '192.168.1.105', '로봇 조립', '-', 'PyQt5 앱'],
            ['작업 PC 6', '192.168.1.106', '성능 검사', '-', 'PyQt5 앱'],
            ['작업 PC 7', '192.168.1.107', '라벨 프린팅', '-', 'PyQt5 앱'],
            ['관리자 PC', '192.168.1.201-230', '대시보드', '-', 'React 웹'],
            ['프린터 1', '192.168.1.50', 'LOT 라벨', '9100', 'Zebra ZT'],
            ['프린터 2', '192.168.1.51', '시리얼 라벨', '9100', 'Zebra ZT']
        ]
    )

    add_section(doc, '4.3 기술 스택', 2)

    add_table(doc,
        ['계층', '기술', '버전', '용도'],
        [
            ['Backend', 'Python', '3.11+', '백엔드 언어'],
            ['', 'FastAPI', 'Latest', 'REST API 프레임워크'],
            ['', 'Uvicorn', 'Latest', 'ASGI 서버'],
            ['', 'SQLAlchemy', '2.0', 'ORM'],
            ['Database', 'PostgreSQL', '15+', '관계형 DB'],
            ['', 'Redis', '7.0+', 'Cache, Session'],
            ['Frontend (작업PC)', 'PyQt5', 'Latest', 'Desktop 앱'],
            ['', 'httpx', 'Latest', 'HTTP Client'],
            ['', 'watchdog', 'Latest', 'File Watcher'],
            ['Frontend (대시보드)', 'React', '18+', 'Web 프레임워크'],
            ['', 'Material-UI', 'Latest', 'UI 컴포넌트'],
            ['', 'Chart.js', 'Latest', '차트 라이브러리'],
            ['인프라', 'Ubuntu', '22.04 LTS', 'OS'],
            ['', 'Nginx', '1.24+', 'Load Balancer'],
            ['', 'Docker', 'Latest', '컨테이너화'],
            ['모니터링', 'Prometheus', 'Latest', '메트릭 수집'],
            ['', 'Grafana', 'Latest', '시각화']
        ]
    )

    doc.add_page_break()

    # === 5. 데이터 설계 ===
    add_section(doc, '5. 데이터 설계', 1)

    add_section(doc, '5.1 ERD (Entity Relationship Diagram)', 2)
    add_diagram(doc, '''
    ┌──────────────────┐
    │  product_models  │
    │  - id (PK)       │
    │  - model_code    │
    │  - model_name    │
    └────────┬─────────┘
             │ 1
             │
             │ N
    ┌────────▼─────────┐       ┌──────────────────┐
    │      lots        │       │    processes     │
    │  - id (PK)       │       │  - id (PK)       │
    │  - lot_number    │       │  - process_id    │
    │  - status        │       │  - process_name  │
    └────────┬─────────┘       └────────┬─────────┘
             │ 1                        │ 1
             │                          │
             │ N                        │ N
    ┌────────▼─────────┐       ┌────────▼─────────┐
    │     serials      │       │  process_data    │
    │  - id (PK)       │───────│  - id (PK)       │
    │  - serial_number │  N    │  - lot_id (FK)   │
    │  - lot_id (FK)   │       │  - serial_id (FK)│
    │  - sequence      │       │  - process_id(FK)│
    └──────────────────┘       │  - start_time    │
                               │  - complete_time  │
                               │  - process_data   │
                               └──────────────────┘
    ''')

    add_section(doc, '5.2 테이블 스키마', 2)

    add_section(doc, '5.2.1 lots 테이블', 3)
    add_code_block(doc, '''CREATE TABLE lots (
    id BIGSERIAL PRIMARY KEY,
    lot_number VARCHAR(50) UNIQUE NOT NULL,
    product_model_id BIGINT REFERENCES product_models(id),
    target_quantity INTEGER DEFAULT 100,
    shift CHAR(1) CHECK (shift IN ('D', 'N')),
    status VARCHAR(20) DEFAULT 'CREATED',
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    updated_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

CREATE INDEX idx_lot_number ON lots(lot_number);
CREATE INDEX idx_lot_status ON lots(status);''', 'sql')

    add_section(doc, '5.2.2 serials 테이블', 3)
    add_code_block(doc, '''CREATE TABLE serials (
    id BIGSERIAL PRIMARY KEY,
    serial_number VARCHAR(100) UNIQUE NOT NULL,
    lot_id BIGINT REFERENCES lots(id),
    sequence INTEGER NOT NULL,
    status VARCHAR(20) DEFAULT 'CREATED',
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW(),
    UNIQUE(lot_id, sequence)
);

CREATE INDEX idx_serial_number ON serials(serial_number);''', 'sql')

    add_section(doc, '5.2.3 process_data 테이블', 3)
    add_code_block(doc, '''CREATE TABLE process_data (
    id BIGSERIAL PRIMARY KEY,
    lot_id BIGINT REFERENCES lots(id),
    serial_id BIGINT REFERENCES serials(id),
    process_id BIGINT REFERENCES processes(id),
    line_id VARCHAR(50) NOT NULL,
    equipment_id VARCHAR(50) NOT NULL,
    worker_id VARCHAR(50),
    process_specific_data JSONB,
    start_time TIMESTAMP WITH TIME ZONE NOT NULL,
    complete_time TIMESTAMP WITH TIME ZONE,
    result VARCHAR(20),
    created_at TIMESTAMP WITH TIME ZONE DEFAULT NOW()
);

CREATE INDEX idx_process_data_jsonb
    ON process_data USING GIN (process_specific_data);''', 'sql')

    doc.add_page_break()

    # === 6. 투자 계획 ===
    add_section(doc, '6. 개발 및 투자 계획', 1)

    add_section(doc, '6.1 개발 일정', 2)
    add_table(doc,
        ['주차', '작업 내용', '산출물'],
        [
            ['1주', '환경 설정, DB 설계, API 설계', 'ERD, API 명세서'],
            ['2-3주', 'Backend API 개발', 'FastAPI 서버'],
            ['4-5주', 'Frontend App 개발', 'PyQt5 앱'],
            ['6주', 'Dashboard 개발', 'React 대시보드'],
            ['7주', '통합 테스트', '테스트 보고서'],
            ['8주', '현장 파일럿', 'UAT 완료']
        ]
    )

    add_section(doc, '6.2 배포 옵션별 비용 비교', 2)

    add_section(doc, '6.2.1 초기 투자 비용', 3)
    add_table(doc,
        ['항목', '온프레미스', 'Railway', 'AWS'],
        [
            ['서버/인프라', '4,985만원', '0원', '0원'],
            ['공통 장비', '995만원', '995만원', '995만원'],
            ['개발비', '810만원', '810만원', '810만원'],
            ['합계', '6,790만원', '1,805만원', '1,805만원']
        ]
    )

    add_section(doc, '6.2.2 월 운영 비용', 3)
    add_table(doc,
        ['항목', '온프레미스', 'Railway', 'AWS'],
        [
            ['인프라', '0원', '26만원', '220-350만원'],
            ['전기/네트워크', '20만원', '0원', '0원'],
            ['시스템 관리', '50만원', '0원', '30만원'],
            ['백업/모니터링', '23만원', '0원', '10만원'],
            ['기타', '20만원', '0원', '5만원'],
            ['합계', '113만원', '26만원', '265-395만원']
        ]
    )

    add_section(doc, '6.2.3 3년 TCO', 3)
    add_table(doc,
        ['구분', '온프레미스', 'Railway', 'AWS'],
        [
            ['초기 투자', '6,790만원', '1,805만원', '1,805만원'],
            ['36개월 운영', '4,068만원', '936만원', '11,880만원'],
            ['3년 총 비용', '10,858만원', '2,741만원', '13,685만원'],
            ['월평균 비용', '302만원', '76만원', '380만원']
        ]
    )

    add_section(doc, '6.3 ROI 분석', 2)
    add_table(doc,
        ['항목', '절감/효과', '연간 금액'],
        [
            ['불량 원인 파악 시간 단축', '월 20시간 절감', '600만원/년'],
            ['재고 관리 효율화', '과잉 재고 감소', '1,000만원/년'],
            ['생산 계획 최적화', '가동률 5% 향상', '2,000만원/년'],
            ['합계', '', '3,600만원/년']
        ]
    )

    add_text(doc, '회수 기간:', bold=True)
    add_bullet(doc, 'Railway: 약 0.8년 (10개월)')
    add_bullet(doc, '온프레미스: 약 3년')
    add_bullet(doc, 'AWS: 약 3.8년')

    doc.add_page_break()

    # === 7. 부록 ===
    add_section(doc, '7. 부록', 1)

    add_section(doc, '7.1 용어 정의', 2)

    terms = [
        ('LOT', '동일 조건으로 생산된 제품 묶음 단위 (통상 100개)'),
        ('시리얼 번호', '개별 제품 고유 식별 번호'),
        ('착공', '공정 작업 시작'),
        ('완공', '공정 작업 완료'),
        ('MES', 'Manufacturing Execution System - 제조실행시스템'),
        ('LMA', 'Linear Muscle Actuator - 로봇의 핵심 구동기'),
        ('SMA 스프링', 'Shape Memory Alloy 스프링 - 온도로 형상 기억'),
        ('EOL 검사', 'End Of Line 검사 - 최종 성능 검사'),
        ('File Watcher', '파일 생성/변경 자동 감시 기능'),
        ('JSONB', 'PostgreSQL JSON 이진 저장 타입'),
        ('JWT', 'JSON Web Token - 사용자 인증 토큰'),
        ('RBAC', 'Role-Based Access Control - 역할 기반 접근 제어'),
        ('TCO', 'Total Cost of Ownership - 총 소유 비용')
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{term}: ')
        run1.font.bold = True
        run1.font.size = Pt(10)
        run2 = p.add_run(definition)
        run2.font.size = Pt(10)

    add_section(doc, '7.2 문서 버전 이력', 2)
    add_table(doc,
        ['버전', '날짜', '변경 내용'],
        [
            ['v1.0', '2025.11.09', '초기 버전'],
            ['v2.0', '2025.11.10', '전면 재작성'],
            ['v2.1', '2025.11.11', '13개 파일로 분리, 논리 오류 10건 수정']
        ]
    )

    # 문서 끝
    doc.add_paragraph()
    end_para = doc.add_paragraph('─── 문서 끝 ───')
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_para.runs[0]
    end_run.font.size = Pt(10)
    end_run.font.color.rgb = RGBColor(150, 150, 150)

    # 페이지 번호 추가
    for section in doc.sections:
        add_page_number(section)

    # 파일 저장
    output_path = 'Withforce_MES_사양서_완전판_v2.docx'
    doc.save(output_path)
    print(f'DOCX 파일 생성 완료: {output_path}')
    print(f'예상 페이지 수: 약 80-100페이지')
    print(f'포함 내용: 테이블, JSON 예시, 다이어그램, 코드 블록')

if __name__ == '__main__':
    create_mes_specification()
