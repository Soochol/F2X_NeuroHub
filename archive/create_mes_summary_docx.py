"""
MES 사양서 요약본 DOCX 생성 스크립트
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(section):
    """페이지 번호 추가"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_mes_specification():
    doc = Document()

    # 페이지 여백 설정
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 제목
    title = doc.add_heading('MES 사양서', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # 부제
    subtitle = doc.add_paragraph('Withforce 웨어러블 로봇 제조실행시스템')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # 빈 줄

    # 문서 정보
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = info.add_run('작성일: 2025-11-11\n문서 유형: 시스템 요구사항 명세서')
    info_text.font.size = Pt(10)
    info_text.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_page_break()

    # 1. 프로젝트 배경 및 목적
    doc.add_heading('1. 프로젝트 배경 및 목적', 1)

    doc.add_heading('1.1 프로젝트 배경', 2)
    doc.add_paragraph(
        'Withforce는 산업용/농업용 허리 보조 로봇을 생산하고 있으며, 현재 생산 공정 전반에 걸친 '
        '데이터 수집 및 추적 체계가 미흡한 상황입니다. 제품 불량 발생 시 원인 파악이 어렵고, '
        '실시간 생산 현황 모니터링이 불가능하여 효율적인 생산 관리에 어려움을 겪고 있습니다.'
    )

    doc.add_heading('1.2 프로젝트 목적', 2)
    p = doc.add_paragraph()
    p.add_run('MES(제조실행시스템) 도입을 통해 다음을 달성합니다:').bold = True

    objectives = [
        '완벽한 제품 추적성 확보 (LOT 및 시리얼 번호 기반)',
        '실시간 생산 현황 모니터링 체계 구축',
        '데이터 기반 품질 관리 및 지속적 개선',
        '공정별 병목 구간 식별 및 생산 최적화'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    # 2. 현재 문제점 (AS-IS)
    doc.add_heading('2. 현재 문제점 (AS-IS)', 1)

    problems = [
        ('추적성 부재', 'LOT 번호 수기 작성, 불량 원인 파악 불가, 시리얼 번호 관리 미흡'),
        ('데이터 수집 미흡', '공정 데이터 미수집, 품질 데이터 분석 불가'),
        ('생산 가시성 부족', '실시간 현황 파악 불가, 병목 구간 식별 어려움'),
        ('수작업 의존', 'LOT 생성 5분 소요, 라벨 수기 작성, 집계 작업 1시간 이상')
    ]

    for title, content in problems:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(content)

    # 3. 비즈니스 요구사항
    doc.add_heading('3. 비즈니스 요구사항', 1)

    doc.add_heading('3.1 핵심 목표', 2)
    goals = [
        'LOT 생성 시간: 5분 → 30초 (90% 단축)',
        '불량 원인 파악: 불가능 → 5분 이내 추적 가능',
        '생산 현황 파악: 1시간 수동 집계 → 실시간 자동 집계',
        '품질 데이터 수집: 미수집 → 100% 자동 수집'
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')

    doc.add_heading('3.2 사용자 정의', 2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'

    headers = ['사용자 유형', '역할', '주요 업무']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    users = [
        ('생산 관리자', 'LOT 생성 및 모니터링', '생산 계획, 현황 모니터링, 통계 분석'),
        ('현장 작업자', '공정 작업 수행', '착공/완공 처리, 데이터 입력'),
        ('시스템 관리자', '시스템 운영 관리', '백업, 사용자 관리, 장애 대응')
    ]

    for i, (user_type, role, tasks) in enumerate(users, 1):
        table.rows[i].cells[0].text = user_type
        table.rows[i].cells[1].text = role
        table.rows[i].cells[2].text = tasks

    doc.add_paragraph()

    # 4. 기능 요구사항
    doc.add_heading('4. 기능 요구사항', 1)

    doc.add_heading('4.1 LOT 관리', 2)
    lot_requirements = [
        'LOT 번호 자동 발급 (형식: WF-KR-YYMMDDX-nnn)',
        'LOT 상태 관리 (생성 → 진행중 → 완료 → 종료)',
        'LOT 바코드 라벨 자동 출력 (Zebra 프린터 연동)',
        'LOT당 생산 목표: 100개',
        '일일 생산 규모: 약 100대'
    ]
    for req in lot_requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('4.2 시리얼 번호 관리', 2)
    serial_requirements = [
        '개별 제품별 고유 시리얼 번호 자동 생성',
        '형식: WF-KR-YYMMDDX-nnn-YYYY',
        '시리얼 번호 바코드 라벨 자동 출력',
        '라벨 재출력 기능 (손상/분실 시)',
        '전체 생산 이력 추적 가능'
    ]
    for req in serial_requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('4.3 공정 관리', 2)
    doc.add_paragraph('8개 공정별 착공/완공 관리가 필요합니다:')

    processes = [
        '레이저 마킹 (Laser Marking)',
        'LMA 조립 (LMA Assembly)',
        'TOF 센서 검사 (TOF Sensor Test)',
        '펌웨어 업로드 (Firmware Upload)',
        '로봇 조립 (Robot Assembly)',
        '성능 검사 (Performance Test)',
        '라벨 프린팅 (Label Printing)',
        '포장 및 외관 검사 (Packing & Visual Inspection)'
    ]

    for process in processes:
        doc.add_paragraph(process, style='List Number')

    doc.add_paragraph()
    doc.add_paragraph('공정 관리 요구사항:')
    process_requirements = [
        '바코드 스캔 기반 착공 등록',
        '공정 순서 제어 (이전 공정 미완료 시 착공 불가)',
        '완공 데이터 자동 수집 (JSON 파일 기반)',
        '공정별 측정 데이터 저장 (온도, 변위, 힘, 부품 LOT 등)',
        '합격/불합격 자동 판정',
        '작업자 및 작업 시간 자동 기록'
    ]
    for req in process_requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_page_break()

    doc.add_heading('4.4 실시간 모니터링 대시보드', 2)
    monitoring_requirements = [
        '금일 생산 현황 (착공/완공/불량 건수)',
        'LOT별 진행 상태 실시간 표시',
        '공정별 현황 및 병목 구간 시각화',
        'LOT 상세 조회 (공정 이력, 시리얼 목록)',
        '시리얼 번호 기반 전체 이력 추적',
        '불량 현황 및 통계 자동 집계'
    ]
    for req in monitoring_requirements:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('4.5 품질 관리', 2)
    quality_requirements = [
        '불량 유형 코드 체계 (8개 대분류)',
        '공정별 불량 원인 추적',
        '불량률 통계 및 트렌드 분석',
        '재작업(Rework) 이력 관리',
        '품질 데이터 기반 지속적 개선'
    ]
    for req in quality_requirements:
        doc.add_paragraph(req, style='List Bullet')

    # 5. 데이터 요구사항
    doc.add_heading('5. 데이터 요구사항', 1)

    doc.add_heading('5.1 핵심 데이터 테이블', 2)

    tables_info = [
        ('LOT 테이블', 'LOT 번호, 상태, 목표 수량, 실제 생산량, 작업대차 번호'),
        ('시리얼 번호 테이블', '시리얼 번호, LOT 번호, 상태, 라벨 출력 이력'),
        ('공정 작업 테이블', '작업 ID, 공정 유형, LOT 번호, 시리얼 번호, 착공/완공 시각, 작업자, 측정 데이터'),
        ('불량 이력 테이블', '불량 ID, 시리얼 번호, 공정, 불량 유형, 조치 내역'),
        ('부품 투입 테이블', '투입 ID, 공정, 부품 LOT 번호, 투입 시각')
    ]

    for table_name, fields in tables_info:
        p = doc.add_paragraph()
        p.add_run(f'{table_name}: ').bold = True
        p.add_run(fields)

    doc.add_heading('5.2 LOT 번호 체계', 2)
    doc.add_paragraph('형식: WF-KR-YYMMDDX-nnn')
    doc.add_paragraph('예시: WF-KR-251110A-001 (2025년 11월 10일 A교대 1번째 LOT)')

    doc.add_heading('5.3 시리얼 번호 체계', 2)
    doc.add_paragraph('형식: WF-KR-YYMMDDX-nnn-YYYY')
    doc.add_paragraph('예시: WF-KR-251110A-001-0042 (해당 LOT의 42번째 제품)')

    doc.add_heading('5.4 불량 코드 체계', 2)
    doc.add_paragraph('8개 대분류:')
    defect_codes = [
        'M: Material (재료 불량)',
        'A: Assembly (조립 불량)',
        'E: Electrical (전기/전자 불량)',
        'P: Performance (성능 불량)',
        'V: Visual (외관 불량)',
        'S: Software (소프트웨어 불량)',
        'D: Dimensional (치수 불량)',
        'O: Other (기타 불량)'
    ]
    for code in defect_codes:
        doc.add_paragraph(code, style='List Bullet')

    doc.add_heading('5.5 공정별 수집 데이터', 2)
    doc.add_paragraph(
        '각 공정별로 특정 측정 데이터를 자동 수집하여 저장해야 합니다. '
        '예: 레이저 마킹(온도, 전력), LMA 조립(토크, 체결 시간), '
        'TOF 센서 검사(측정값), 성능 검사(온도, 변위, 힘) 등.'
    )

    # 6. 성능 요구사항
    doc.add_heading('6. 성능 요구사항', 1)

    doc.add_heading('6.1 응답 시간', 2)
    performance = [
        'LOT 생성: 30초 이내',
        '착공 API 응답: 1초 이내',
        '완공 데이터 처리: 2초 이내',
        '대시보드 로딩: 3초 이내',
        '시리얼 추적 조회: 5초 이내'
    ]
    for perf in performance:
        doc.add_paragraph(perf, style='List Bullet')

    doc.add_heading('6.2 동시성 요구사항', 2)
    concurrency = [
        '동시 접속자: 100명 (작업자 70명 + 관리자 30명)',
        '일일 트랜잭션: 약 50,000건',
        '동시 착공/완공 처리: 20건/초',
        '피크 시간대 처리 능력 확보'
    ]
    for conc in concurrency:
        doc.add_paragraph(conc, style='List Bullet')

    doc.add_heading('6.3 가용성 및 안정성', 2)
    availability = [
        '시스템 가동률: 99% 이상 (월 7.2시간 이내 다운타임)',
        '데이터 백업: 일 1회 자동 백업',
        '데이터 보관 기간: 최소 3년',
        '장애 복구 시간: 1시간 이내'
    ]
    for avail in availability:
        doc.add_paragraph(avail, style='List Bullet')

    doc.add_heading('6.4 보안 요구사항', 2)
    security = [
        '사용자 인증 및 권한 관리',
        '역할 기반 접근 제어 (RBAC)',
        '데이터 암호화 (전송 중/저장 시)',
        '작업 이력 로그 기록 및 감사',
        '부적절한 접근 시도 차단'
    ]
    for sec in security:
        doc.add_paragraph(sec, style='List Bullet')

    doc.add_page_break()

    # 7. 검수 기준
    doc.add_heading('7. 검수 기준', 1)

    doc.add_paragraph('시스템 구축 완료 후 다음 검수 기준을 충족해야 합니다:')

    doc.add_heading('7.1 기능 검수', 2)

    acceptance_categories = [
        ('LOT 관리', [
            'LOT 번호 자동 발급 정상 작동',
            'LOT 바코드 라벨 출력 성공',
            'LOT 상태 변경 정상 처리',
            'LOT 생성 시간 30초 이내'
        ]),
        ('시리얼 번호 관리', [
            '시리얼 번호 자동 생성 정상 작동',
            '시리얼 라벨 출력 성공',
            '라벨 재출력 기능 정상',
            '전체 이력 추적 가능'
        ]),
        ('공정 관리', [
            '바코드 스캔 기반 착공 등록 정상',
            '공정 순서 제어 정상 작동',
            '완공 데이터 자동 수집 성공',
            '측정 데이터 정확히 저장',
            '합격/불합격 판정 정상'
        ]),
        ('모니터링', [
            '실시간 생산 현황 정확히 표시',
            'LOT 진행 상태 실시간 업데이트',
            '공정별 병목 구간 시각화',
            '불량 통계 정확히 집계'
        ]),
        ('추적성', [
            '시리얼 번호로 전체 이력 조회 가능',
            '부품 LOT 추적 정상',
            '작업자 이력 조회 가능',
            '불량 원인 추적 5분 이내'
        ])
    ]

    for category, items in acceptance_categories:
        p = doc.add_paragraph()
        p.add_run(f'{category}:').bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('7.2 성능 검수', 2)
    performance_acceptance = [
        '동시 접속자 20명 부하 테스트 통과',
        '응답 시간 요구사항 충족',
        '피크 시간대 안정성 검증',
        '24시간 연속 운영 안정성 확인'
    ]
    for perf in performance_acceptance:
        doc.add_paragraph(perf, style='List Bullet')

    doc.add_heading('7.3 보안 검수', 2)
    security_acceptance = [
        '사용자 인증 및 권한 관리 정상 작동',
        '역할별 접근 제어 확인',
        '데이터 암호화 확인',
        '로그 기록 및 감사 추적 정상'
    ]
    for sec in security_acceptance:
        doc.add_paragraph(sec, style='List Bullet')

    # 8. 부록
    doc.add_heading('8. 부록', 1)

    doc.add_heading('8.1 주요 용어 정의', 2)

    terms = [
        ('LOT', '동일 조건으로 생산된 제품 묶음 단위 (통상 100개)'),
        ('시리얼 번호', '개별 제품 고유 식별 번호'),
        ('착공', '공정 작업 시작'),
        ('완공', '공정 작업 완료'),
        ('Traceability', '추적성 - 제품 이력 추적 능력'),
        ('MES', 'Manufacturing Execution System - 제조실행시스템'),
        ('LMA', 'Linear Muscle Actuator - 로봇의 핵심 구동기'),
        ('SMA 스프링', 'Shape Memory Alloy 스프링 - 온도로 형상 기억하여 구동'),
        ('EOL 검사', 'End Of Line 검사 - 최종 성능 검사'),
        ('TOF 센서', 'Time Of Flight 센서 - 거리 측정 센서')
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        p.add_run(f'{term}: ').bold = True
        p.add_run(definition)

    doc.add_heading('8.2 프로젝트 성공 요소', 2)
    success_factors = [
        '현장 작업자와의 충분한 사전 협의 및 교육',
        '단계적 구축 및 점진적 기능 확대',
        '안정적인 바코드 스캐너 및 프린터 하드웨어 확보',
        '충분한 테스트 기간 확보 (최소 2주)',
        '지속적인 피드백 수렴 및 개선'
    ]
    for factor in success_factors:
        doc.add_paragraph(factor, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph()

    # 문서 끝
    end_para = doc.add_paragraph('문서 끝')
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_run = end_para.runs[0]
    end_run.font.size = Pt(10)
    end_run.font.color.rgb = RGBColor(150, 150, 150)

    # 페이지 번호 추가
    for section in doc.sections:
        add_page_number(section)

    # 파일 저장
    output_path = 'MES_사양서_요약본.docx'
    doc.save(output_path)
    print(f'DOCX 파일 생성 완료: {output_path}')
    print(f'예상 페이지 수: 약 5-6페이지')

if __name__ == '__main__':
    create_mes_specification()
