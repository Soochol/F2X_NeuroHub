#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
F2X NeuroHub MES 사양서 - Markdown to Professional DOCX 변환

Microsoft 공식 문서 스타일 적용
"""

import re
import base64
import urllib.parse
import urllib.request
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def setup_styles(doc):
    """Modern Tech Datasheet 스타일 설정"""

    # Normal 스타일
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)  # 약간 작게 (datasheet 스타일)

    # 한글 폰트
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(8)
    paragraph_format.line_spacing = 1.2  # 더 타이트한 간격

    # Heading 1 - Modern Tech Style (Boxed)
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Calibri'
    heading1.font.size = Pt(20)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(255, 255, 255)  # White text
    rFonts = heading1.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    heading1.paragraph_format.space_before = Pt(18)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.keep_with_next = True

    # Heading 2 - Blue box with white text
    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Calibri'
    heading2.font.size = Pt(16)
    heading2.font.bold = True
    heading2.font.color.rgb = RGBColor(255, 255, 255)  # White text
    rFonts = heading2.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    heading2.paragraph_format.space_before = Pt(14)
    heading2.paragraph_format.space_after = Pt(10)
    heading2.paragraph_format.keep_with_next = True

    # Heading 3 - Gray background
    heading3 = doc.styles['Heading 3']
    heading3.font.name = 'Calibri'
    heading3.font.size = Pt(13)
    heading3.font.bold = True
    heading3.font.color.rgb = RGBColor(31, 78, 120)  # Dark blue text
    rFonts = heading3.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)
    heading3.paragraph_format.keep_with_next = True

    # Heading 4 - Simple bold
    heading4 = doc.styles['Heading 4']
    heading4.font.name = 'Calibri'
    heading4.font.size = Pt(11)
    heading4.font.bold = True
    heading4.font.color.rgb = RGBColor(68, 68, 68)
    rFonts = heading4.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    heading4.paragraph_format.space_before = Pt(10)
    heading4.paragraph_format.space_after = Pt(4)


def set_cell_background(cell, fill):
    """테이블 셀 배경색 설정"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)


def add_boxed_heading(doc, text, level):
    """Premium Modern Tech Style: 세련된 박스형 제목"""

    # 색상 정의 (더 세련된 색상 팔레트)
    colors = {
        1: ('0F2C4C', RGBColor(255, 255, 255)),  # Deep navy bg, white text
        2: ('1F4E78', RGBColor(255, 255, 255)),  # Dark blue bg, white text
        3: ('F5F5F5', RGBColor(31, 78, 120)),    # Very light gray bg, dark blue text
        4: (None, RGBColor(31, 78, 120))         # No background, dark blue text
    }

    bg_color, text_color = colors.get(level, colors[4])

    para = doc.add_paragraph()
    para.paragraph_format.keep_with_next = True

    # 레벨별 스타일 적용
    if level == 1:
        para.paragraph_format.space_before = Pt(20)
        para.paragraph_format.space_after = Pt(14)
        font_size = Pt(22)
        left_padding = Pt(15)
        vertical_padding = '120'  # 6pt
    elif level == 2:
        para.paragraph_format.space_before = Pt(16)
        para.paragraph_format.space_after = Pt(12)
        font_size = Pt(17)
        left_padding = Pt(12)
        vertical_padding = '100'  # 5pt
    elif level == 3:
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(8)
        font_size = Pt(13.5)
        left_padding = Pt(10)
        vertical_padding = '80'   # 4pt
    else:
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        font_size = Pt(11.5)
        left_padding = Pt(0)
        vertical_padding = '0'

    # 배경색 및 스타일링 (Level 1-3)
    if bg_color:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg_color)
        para._element.get_or_add_pPr().append(shading)

        # 패딩 추가 (상하좌우)
        para.paragraph_format.left_indent = left_padding
        para.paragraph_format.right_indent = Pt(10)

        # 내부 간격
        pPr = para._element.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), vertical_padding)
        spacing.set(qn('w:after'), vertical_padding)
        pPr.append(spacing)

        # Level 1, 2에 왼쪽 accent 바 추가
        if level in [1, 2]:
            pBdr = OxmlElement('w:pBdr')
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), '48')  # 6pt 두께
            left_border.set(qn('w:space'), '4')
            accent_color = '4A90E2' if level == 1 else '5BA3F5'  # Bright blue accent
            left_border.set(qn('w:color'), accent_color)
            pBdr.append(left_border)
            pPr.append(pBdr)

    # 텍스트 추가
    run = para.add_run(text)
    run.font.name = 'Calibri Light' if level in [1, 2] else 'Calibri'
    run.font.size = font_size
    run.font.bold = True
    run.font.color.rgb = text_color

    # 한글 폰트
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    rPr.append(rFonts)


def add_page_number(section):
    """페이지 번호 추가"""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 좌측에 문서명
    run = paragraph.add_run('Withforce MES 사양서 v2.0')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 탭
    paragraph.add_run('\t\t')

    # 페이지 번호
    run = paragraph.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def create_cover_page(doc):
    """Premium Modern Tech 표지 - 세련되고 고급스러운 디자인"""

    # 상단 여백
    for _ in range(3):
        doc.add_paragraph()

    # 회사명 (작게, 대문자, 간격있는 레터링)
    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = company.add_run('W I T H F O R C E')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 120, 120)
    run.font.name = 'Calibri Light'
    run.font.bold = False

    # 레터 스페이싱
    rPr = run._element.get_or_add_rPr()
    spacing_elem = OxmlElement('w:spacing')
    spacing_elem.set(qn('w:val'), '80')  # 문자 간격
    rPr.append(spacing_elem)

    # 얇은 accent 라인
    accent_line = doc.add_paragraph()
    accent_line.paragraph_format.space_before = Pt(6)
    accent_line.paragraph_format.space_after = Pt(20)

    pBdr = OxmlElement('w:pBdr')
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '18')  # 2.25pt
    top_border.set(qn('w:space'), '1')
    top_border.set(qn('w:color'), '4A90E2')  # Bright blue accent
    pBdr.append(top_border)
    accent_line._element.get_or_add_pPr().append(pBdr)

    # 메인 타이틀 박스 (Deep navy background with accent)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    title_para.paragraph_format.left_indent = Pt(20)
    title_para.paragraph_format.right_indent = Pt(15)

    # 배경색
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '0F2C4C')  # Deep navy
    title_para._element.get_or_add_pPr().append(shading)

    # 내부 간격 증가
    pPr = title_para._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')  # 12pt
    spacing.set(qn('w:after'), '240')   # 12pt
    pPr.append(spacing)

    # 왼쪽 accent 바
    pBdr = OxmlElement('w:pBdr')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '60')  # 7.5pt 두께
    left_border.set(qn('w:space'), '4')
    left_border.set(qn('w:color'), '4A90E2')  # Bright blue accent
    pBdr.append(left_border)
    pPr.append(pBdr)

    run = title_para.add_run('MES 시스템 개선안')
    run.font.size = Pt(44)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = 'Calibri Light'

    # 버전 박스 (더 부드러운 블루)
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    version_para.paragraph_format.space_after = Pt(0)
    version_para.paragraph_format.left_indent = Pt(20)
    version_para.paragraph_format.right_indent = Pt(15)

    # 배경색
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '1F4E78')  # Dark blue
    version_para._element.get_or_add_pPr().append(shading)

    # 내부 간격
    pPr = version_para._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '120')
    spacing.set(qn('w:after'), '120')
    pPr.append(spacing)

    run = version_para.add_run('Version 2.0')
    run.font.size = Pt(19)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.name = 'Calibri'

    # 공간
    for _ in range(3):
        doc.add_paragraph()

    # 문서 타입 (미니멀한 스타일)
    doc_type = doc.add_paragraph()
    doc_type.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc_type.paragraph_format.left_indent = Pt(2)

    run = doc_type.add_run('시스템 요구사항 명세서')
    run.font.size = Pt(16)
    run.font.bold = False
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = 'Calibri Light'

    # 하단 정보
    for _ in range(10):
        doc.add_paragraph()

    # 메타 정보 테이블 (미니멀한 디자인)
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'

    info_data = [
        ('작성일', '2025년 11월 10일'),
        ('작성자', 'Claude'),
        ('문서 분류', '기술 사양서')
    ]

    for idx, (label, value) in enumerate(info_data):
        row = info_table.rows[idx]

        # 레이블 셀
        label_cell = row.cells[0]
        label_cell.text = label
        set_cell_background(label_cell, 'F8F9FA')
        label_cell.width = Inches(1.5)

        for para in label_cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(90, 90, 90)
                run.font.name = 'Calibri'

        # 값 셀
        value_cell = row.cells[1]
        value_cell.text = value

        for para in value_cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(68, 68, 68)
                run.font.name = 'Calibri'

    # 테이블 테두리 (얇고 부드럽게)
    tbl = info_table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '2')  # 매우 얇게
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'E8E8E8')
        tblBorders.append(border)

    # 수평선만 약간 두껍게
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'E8E8E8')
    tblBorders.append(insideH)

    tblPr.append(tblBorders)

    doc.add_page_break()


def create_toc_placeholder(doc):
    """목차 페이지"""
    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run('목차')
    run.font.size = Pt(24)
    run.font.bold = False
    run.font.color.rgb = RGBColor(46, 116, 181)
    run.font.name = 'Calibri Light'

    doc.add_paragraph()

    # 안내
    toc_note = doc.add_paragraph()
    run = toc_note.add_run('Microsoft Word에서 참조 > 목차 업데이트를 클릭하여 자동 생성하세요.')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()
    doc.add_page_break()


def apply_inline_formatting(paragraph, text):
    """인라인 포맷팅"""
    pattern = r'(\*\*.*?\*\*|`.*?`|\*.*?\*)'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if re.match(r'^\*\*.*\*\*$', part):
            # 굵게
            clean_text = part.strip('*')
            run = paragraph.add_run(clean_text)
            run.bold = True
        elif re.match(r'^`.*`$', part):
            # 인라인 코드
            clean_text = part.strip('`')
            run = paragraph.add_run(clean_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(231, 76, 60)  # Red
            # 배경
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'F0F0F0')
            run._element.get_or_add_rPr().append(shading)
        elif re.match(r'^\*.*\*$', part):
            # 이탤릭
            clean_text = part.strip('*')
            run = paragraph.add_run(clean_text)
            run.italic = True
        else:
            paragraph.add_run(part)


def create_styled_table(doc, headers, rows):
    """Premium Modern Tech 스타일 테이블"""
    table = doc.add_table(rows=1, cols=len(headers))

    # 헤더 행 - Deep navy with premium feel
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header.strip('*').strip()

        # Deep navy header
        set_cell_background(cell, '0F2C4C')

        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Calibri'

        cell.vertical_alignment = 1

        # 셀 패딩 증가 (더 여유로운 느낌)
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'bottom']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '100')  # 5pt
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        for margin_name in ['left', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '120')  # 6pt
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)

    # 데이터 행
    for idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            if i < len(row_cells):
                cell = row_cells[i]
                cell.text = cell_data.strip()

                # 셀 스타일
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(4)
                    paragraph.paragraph_format.space_after = Pt(4)
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri'
                        run.font.color.rgb = RGBColor(51, 51, 51)

                # 셀 패딩 증가
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcMar = OxmlElement('w:tcMar')
                for margin_name in ['top', 'bottom']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '80')  # 4pt
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                for margin_name in ['left', 'right']:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), '120')  # 6pt
                    margin.set(qn('w:type'), 'dxa')
                    tcMar.append(margin)
                tcPr.append(tcMar)

                # 교차 행 배경 (더 부드러운 색상)
                if idx % 2 == 1:
                    set_cell_background(cell, 'F8F9FA')  # Very light gray
                else:
                    set_cell_background(cell, 'FFFFFF')  # White

    # 테이블 스타일 - 깔끔한 선
    table.style = 'Table Grid'

    # 테이블 테두리 커스터마이징 (더 세련된 스타일)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 테두리 설정 (헤더는 없고, 데이터 행은 얇은 회색)
    tblBorders = OxmlElement('w:tblBorders')

    # 외부 테두리 없음
    for border_name in ['top', 'left', 'right', 'bottom']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')  # 0.75pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'E0E0E0')  # Very light gray
        tblBorders.append(border)

    # 내부 수평선 (얇고 부드러운)
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')  # 0.5pt
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'E8E8E8')
    tblBorders.append(insideH)

    # 내부 수직선 (더 얇게)
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'single')
    insideV.set(qn('w:sz'), '2')  # 0.25pt
    insideV.set(qn('w:space'), '0')
    insideV.set(qn('w:color'), 'F0F0F0')
    tblBorders.append(insideV)

    tblPr.append(tblBorders)

    # 테이블 앞뒤 여백 증가
    table_para = doc.paragraphs[-1] if doc.paragraphs else None
    if table_para:
        table_para.paragraph_format.space_after = Pt(16)

    return table


def add_code_block(doc, code_lines, language=''):
    """코드 블록"""

    # 언어 표시
    if language and language not in ['mermaid']:
        lang_para = doc.add_paragraph()
        run = lang_para.add_run(f'  {language.upper()}')
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        lang_para.paragraph_format.space_after = Pt(2)
        # 코드 블록이 제목과 분리되지 않도록
        lang_para.paragraph_format.keep_with_next = True

    # 코드
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Inches(0.25)
    code_para.paragraph_format.space_before = Pt(6)
    code_para.paragraph_format.space_after = Pt(12)
    # 코드 블록 내부에서 페이지 나누기 최소화
    code_para.paragraph_format.keep_together = True

    # 배경색 (연한 회색)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    code_para._element.get_or_add_pPr().append(shading)

    # 테두리 (왼쪽에만 파란색 막대)
    pBdr = OxmlElement('w:pBdr')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '24')  # 3pt
    left_border.set(qn('w:space'), '4')
    left_border.set(qn('w:color'), '2E74B5')
    pBdr.append(left_border)
    code_para._element.get_or_add_pPr().append(pBdr)

    code_text = '\n'.join(code_lines)
    run = code_para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(51, 51, 51)


def add_callout_box(doc, text, box_type='info'):
    """Microsoft 스타일 콜아웃 박스"""

    # 색상 설정
    colors = {
        'info': ('D6ECF7', '2E74B5'),      # Blue
        'warning': ('FFF3CD', 'FFC107'),   # Yellow
        'success': ('D4EDDA', '28A745'),   # Green
        'danger': ('F8D7DA', 'DC3545')     # Red
    }

    bg_color, border_color = colors.get(box_type, colors['info'])

    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.right_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)

    # 배경색
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), bg_color)
    para._element.get_or_add_pPr().append(shading)

    # 왼쪽 테두리
    pBdr = OxmlElement('w:pBdr')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '24')
    left_border.set(qn('w:space'), '4')
    left_border.set(qn('w:color'), border_color)
    pBdr.append(left_border)
    para._element.get_or_add_pPr().append(pBdr)

    run = para.add_run('  ' + text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'


def mermaid_to_image(mermaid_code):
    """Mermaid 코드를 이미지로 변환 (mermaid.ink API 사용)"""
    try:
        # Mermaid 코드를 base64로 인코딩
        encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')

        # mermaid.ink API URL
        url = f"https://mermaid.ink/img/{encoded}"

        # 이미지 다운로드
        with urllib.request.urlopen(url, timeout=10) as response:
            image_data = response.read()
            return BytesIO(image_data)
    except Exception as e:
        print(f"[WARNING] Mermaid 다이어그램 변환 실패: {e}")
        return None


def add_process_flow_diagram(doc):
    """2.3절 현재 생산 공정 흐름을 도형 객체로 생성 (개선된 디자인)"""

    # 공정 단계 정의 (번호 포함)
    processes = [
        ("1", "레이저 마킹"),
        ("2", "LMA 조립"),
        ("3", "센서 검사"),
        ("4", "펌웨어 업로드"),
        ("5", "로봇 조립"),
        ("6", "성능검사"),
        ("7", "라벨 프린팅"),
        ("8", "포장 + 외관검사")
    ]

    # 색상 팔레트 (그라데이션 효과)
    colors = [
        '1F4E78',  # 진한 파랑
        '2E5C8A',  # 파랑
        '3D6A9C',  # 중간 파랑
        '4C78AE',  # 밝은 파랑
        '5B86C0',  # 더 밝은 파랑
        '6A94D2',  # 연한 파랑
        '79A2E4',  # 아주 연한 파랑
        '88B0F6',  # 하늘색
    ]

    # 다이어그램 제목을 위한 paragraph (keep with next)
    diagram_intro = doc.add_paragraph()
    diagram_intro.paragraph_format.keep_with_next = True
    diagram_intro.paragraph_format.space_before = Pt(12)
    diagram_intro.paragraph_format.space_after = Pt(12)

    # 2행 구조: 공정 박스, 화살표
    table = doc.add_table(rows=2, cols=len(processes))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 각 공정 블록 스타일링
    for i, (num, process) in enumerate(processes):
        cell = table.rows[0].cells[i]

        # 배경색 (그라데이션)
        set_cell_background(cell, colors[i])

        # 셀 내용: 번호 + 공정명 (두 줄)
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 번호 (작고 연한 텍스트)
        num_run = para.add_run(f"{num}\n")
        num_run.font.size = Pt(8)
        num_run.font.bold = False
        num_run.font.color.rgb = RGBColor(255, 255, 255)
        num_run.font.name = 'Calibri'

        # 공정명 (크고 굵은 텍스트)
        process_run = para.add_run(process)
        process_run.font.size = Pt(10)
        process_run.font.bold = True
        process_run.font.color.rgb = RGBColor(255, 255, 255)
        process_run.font.name = 'Malgun Gothic'

        # 셀 크기 및 여백 조정
        cell.width = Inches(0.9)
        cell.vertical_alignment = 1  # 수직 중앙 정렬

        # 셀 패딩 증가
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name in ['top', 'bottom', 'left', 'right']:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), '100')
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)
        tcPr.append(tcMar)

        # 화살표 추가 (마지막 셀 제외)
        if i < len(processes) - 1:
            arrow_cell = table.rows[1].cells[i]
            arrow_cell.text = ""

            arrow_para = arrow_cell.paragraphs[0]
            arrow_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            arrow_run = arrow_para.add_run("▶")
            arrow_run.font.size = Pt(14)
            arrow_run.font.color.rgb = RGBColor(46, 116, 181)
            arrow_run.font.bold = True

            arrow_cell.vertical_alignment = 1

    # 마지막 화살표 셀 비우기
    table.rows[1].cells[-1].text = ""

    # 테이블 테두리 제거 (깔끔한 디자인)
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 테두리 설정 (공정 박스 테두리만)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # 테이블 행이 페이지 경계를 넘어가지 않도록
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

    # 테이블 간격
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(18)

    print("[OK] 공정 흐름도 도형 객체 생성 완료 (개선된 디자인)")


def add_mermaid_diagram(doc, mermaid_code, section_context=''):
    """Mermaid 다이어그램을 이미지로 변환하거나 특수 처리"""

    # 2.3절 공정 흐름도는 도형 객체로 생성
    if '현재 생산 공정 흐름' in section_context or 'graph LR' in mermaid_code:
        add_process_flow_diagram(doc)
        return

    # 이미지 변환 시도
    image_stream = mermaid_to_image(mermaid_code)

    if image_stream:
        # 이미지 삽입
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)

        try:
            run = para.add_run()
            run.add_picture(image_stream, width=Inches(5.5))
            print("[OK] Mermaid 다이어그램 이미지 삽입 성공")
        except Exception as e:
            print(f"[WARNING] 이미지 삽입 실패: {e}")
            # 실패 시 대체 텍스트
            add_callout_box(doc, ' 다이어그램: 원본 마크다운 문서를 참조하세요', 'warning')
    else:
        # API 실패 시 대체 텍스트
        add_callout_box(doc, ' 다이어그램: 원본 마크다운 문서를 참조하세요', 'warning')


def parse_markdown_line(line):
    """마크다운 파싱"""
    if line.startswith('#'):
        level = len(line) - len(line.lstrip('#'))
        text = line.lstrip('#').strip()
        return ('heading', level, text)

    if re.match(r'^\|[\s:-]+\|', line):
        return ('table_separator', None, None)

    if line.startswith('|') and line.endswith('|'):
        cells = [cell.strip() for cell in line.strip('|').split('|')]
        return ('table_row', None, cells)

    if line.startswith('```'):
        lang = line[3:].strip()
        return ('code_fence', None, lang)

    if re.match(r'^\s*[-*+]\s+', line):
        indent = len(line) - len(line.lstrip())
        text = re.sub(r'^\s*[-*+]\s+', '', line)
        return ('bullet_list', indent, text)

    if re.match(r'^\s*\d+\.\s+', line):
        indent = len(line) - len(line.lstrip())
        text = re.sub(r'^\s*\d+\.\s+', '', line)
        return ('numbered_list', indent, text)

    if re.match(r'^[-*_]{3,}$', line.strip()):
        return ('horizontal_rule', None, None)

    if not line.strip():
        return ('empty', None, None)

    return ('text', None, line)


def convert_markdown_to_docx(md_file, docx_file):
    """마크다운을 Microsoft 스타일 DOCX로 변환"""

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # A4
    section.page_width = Inches(8.27)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    setup_styles(doc)
    add_page_number(section)
    create_cover_page(doc)
    # create_toc_placeholder(doc)  # 목차 제거

    # 본문 파싱
    in_code_block = False
    code_lines = []
    code_language = ''
    current_table_headers = None
    current_table_rows = []
    skip_until_heading = False
    current_section = ''  # 현재 섹션 추적

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        if not skip_until_heading:
            if line.startswith('# '):
                skip_until_heading = True
            else:
                i += 1
                continue

        parsed = parse_markdown_line(line)
        ptype, level, content = parsed

        # 코드 블록
        if ptype == 'code_fence':
            if not in_code_block:
                in_code_block = True
                code_language = content or ''
                code_lines = []
            else:
                in_code_block = False

                if code_language.lower() == 'mermaid':
                    # Mermaid 다이어그램 이미지로 변환 (또는 특수 처리)
                    mermaid_code = '\n'.join(code_lines)
                    add_mermaid_diagram(doc, mermaid_code, current_section)
                else:
                    add_code_block(doc, code_lines, code_language)

                code_lines = []
                code_language = ''

            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 테이블 완료
        if current_table_headers and ptype not in ['table_row', 'table_separator']:
            if current_table_rows:
                create_styled_table(doc, current_table_headers, current_table_rows)
                doc.add_paragraph()
            current_table_headers = None
            current_table_rows = []

        # 제목
        if ptype == 'heading':
            if level <= 4:
                # 주요 섹션(Level 1, 2) 시작 시 페이지 나누기
                if level == 1:
                    # 대섹션 시작 전 페이지 나누기 (첫 섹션 제외)
                    if content not in ['프로젝트 개요']:
                        doc.add_page_break()
                elif level == 2:
                    # 중요한 하위 섹션 시작 전 페이지 나누기
                    important_subsections = [
                        '핵심 기능 요구사항',
                        '백엔드 서버',
                        '데이터베이스 설계',
                        '개발 일정'
                    ]
                    if any(keyword in content for keyword in important_subsections):
                        doc.add_page_break()

                # Modern Tech Style: 박스형 제목 사용
                add_boxed_heading(doc, content, level)
                # 현재 섹션 저장 (2.3 현재 생산 공정 흐름 감지용)
                current_section = content

        # 표
        elif ptype == 'table_row':
            if current_table_headers is None:
                current_table_headers = content
            else:
                current_table_rows.append(content)

        elif ptype == 'table_separator':
            pass

        # 리스트
        elif ptype == 'bullet_list':
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, content)

        elif ptype == 'numbered_list':
            p = doc.add_paragraph(style='List Number')
            apply_inline_formatting(p, content)

        # 수평선
        elif ptype == 'horizontal_rule':
            hr = doc.add_paragraph()
            hr.paragraph_format.space_before = Pt(12)
            hr.paragraph_format.space_after = Pt(12)

            # 테두리
            pBdr = OxmlElement('w:pBdr')
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '6')
            bottom_border.set(qn('w:space'), '1')
            bottom_border.set(qn('w:color'), 'D0D0D0')
            pBdr.append(bottom_border)
            hr._element.get_or_add_pPr().append(pBdr)

        # 빈 줄
        elif ptype == 'empty':
            pass

        # 텍스트
        elif ptype == 'text':
            if content.startswith('**문서 버전:**') or content.startswith('**작성일:**') or content.startswith('**문서 유형:**'):
                i += 1
                continue

            if content.startswith('[') and '](' in content:
                i += 1
                continue

            p = doc.add_paragraph()
            apply_inline_formatting(p, content)

        i += 1

    # 마지막 테이블
    if current_table_headers and current_table_rows:
        create_styled_table(doc, current_table_headers, current_table_rows)

    doc.save(docx_file)
    print(f"[OK] Microsoft 스타일 DOCX 생성 완료: {docx_file}")


if __name__ == '__main__':
    md_file = Path(__file__).parent / 'F2X_NeuroHub_MES_사양서_v2.0.md'
    docx_file = Path(__file__).parent / 'Withforce_MES_사양서_v2.0.docx'

    if not md_file.exists():
        print(f"[ERROR] 마크다운 파일을 찾을 수 없습니다: {md_file}")
        exit(1)

    print(f"[INFO] Microsoft 스타일로 변환 시작...")
    convert_markdown_to_docx(md_file, docx_file)
    print(f"[OK] 완료! Word에서 열어보세요.")
